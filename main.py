import streamlit as st
import re
import os
from datetime import datetime
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from collections import defaultdict
from difflib import Differ
import base64
from sklearn.feature_extraction.text import TfidfVectorizer
import spacy
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

nlp = spacy.load("en_core_web_sm")

# --- Constants ---
ACTION_VERBS = [
    "achieved", "managed", "developed", "led", "optimized", 
    "implemented", "increased", "reduced", "designed", "built"
]

QUANTIFIABLE_PATTERNS = [
    r"\d+%", r"\$\d+", r"\d+ (users|customers|projects)",
    r"increased by \d+", r"reduced by \d+",
    r"\b(increased|decreased|boosted|saved|generated)\s+.*?\d+%?"
]

SECTIONS = ["experience", "education", "skills", "projects"]

# --- Core Functions ---
def extract_text_from_pdf(file):
    text = ""
    reader = PdfReader(file)
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def clean_text(text):
    text = re.sub(r'http\S+|www\S+|https\S+', '', text)
    text = re.sub(r'\W', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text.lower()

def extract_keywords_tfidf(text, top_n=10):
    vectorizer = TfidfVectorizer(stop_words='english')
    tfidf = vectorizer.fit_transform([text])
    scores = zip(vectorizer.get_feature_names_out(), tfidf.toarray()[0])
    sorted_scores = sorted(scores, key=lambda x: x[1], reverse=True)
    return [word for word, score in sorted_scores[:top_n]]

def extract_sections(text):
    extracted = {}
    for i, section in enumerate(SECTIONS):
        match = re.search(section, text.lower())
        if match:
            start = match.start()
            end = len(text)
            for s in SECTIONS[i+1:]:
                m = re.search(s, text.lower()[start+1:])
                if m:
                    end = m.start() + start + 1
                    break
            extracted[section] = text[start:end].strip()
    return extracted

def extract_entities(text):
    doc = nlp(text)
    return [(ent.text, ent.label_) for ent in doc.ents if ent.label_ in ("ORG", "PERSON", "GPE", "DATE", "WORK_OF_ART")]

def check_quantifiable_achievements(text):
    sentences = re.split(r'[.!?]', text)
    weak = []
    for sentence in sentences:
        if not any(re.search(p, sentence.lower()) for p in QUANTIFIABLE_PATTERNS):
            if any(verb in sentence.lower() for verb in ACTION_VERBS):
                weak.append(sentence.strip())
    return weak

def analyze_resume_structure(text):
    found = defaultdict(bool)
    for section in SECTIONS:
        found[section] = section in text.lower()
    return found

def generate_diff(old, new):
    d = Differ()
    diff = list(d.compare(old.splitlines(), new.splitlines()))
    return "\n".join(line for line in diff if not line.startswith('? '))

def suggest_improvements(resume_text, job_desc):
    suggestions = []

    resume_kw = set(extract_keywords_tfidf(clean_text(resume_text)))
    job_kw = set(extract_keywords_tfidf(clean_text(job_desc)))
    missing_kw = job_kw - resume_kw
    if missing_kw:
        suggestions.append({
            "category": "Keywords", 
            "suggestion": f"Add missing keywords: {', '.join(missing_kw)}",
            "priority": "High"
        })

    used_verbs = [v for v in ACTION_VERBS if v in resume_text.lower()]
    if len(used_verbs) < 3:
        suggestions.append({
            "category": "Writing Style",
            "suggestion": f"Use more action verbs (only {len(used_verbs)} found). Try: {', '.join(ACTION_VERBS[:3])}",
            "priority": "Medium"
        })

    weak_sentences = check_quantifiable_achievements(resume_text)
    if weak_sentences:
        example = f"{weak_sentences[0]} → 'Increased performance by 30%'"
        suggestions.append({
            "category": "Achievements",
            "suggestion": f"Add metrics: {example}",
            "priority": "High"
        })

    structure = analyze_resume_structure(resume_text)
    missing = [s for s, found in structure.items() if not found]
    if missing:
        suggestions.append({
            "category": "Structure",
            "suggestion": f"Add missing sections: {', '.join(missing)}",
            "priority": "Medium"
        })

    return suggestions

def enhance_docx(docx_file, suggestions):
    doc = Document(docx_file)

    doc.add_heading('Resume Optimization Suggestions', level=1)
    for suggestion in suggestions:
        doc.add_paragraph(f"[{suggestion['priority']}] {suggestion['category']}: {suggestion['suggestion']}")

    optimized_path = f"optimized/optimized_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(optimized_path)
    return optimized_path

def enhance_pdf(pdf_file, suggestions):
    reader = PdfReader(pdf_file)
    writer = PdfWriter()

    # Adding original pages
    for page in reader.pages:
        writer.add_page(page)

    suggestion_text = "\n".join([f"{s['category']} - {s['priority']}:\n{ s['suggestion']}" for s in suggestions])

    # Create a temporary file to save the new PDF with suggestions
    temp_file = "optimized/optimized_resume_with_suggestions.pdf"
    packet = BytesIO()
    c = canvas.Canvas(packet, pagesize=letter)
    c.drawString(72, 750, "Optimized Suggestions:")
    text_object = c.beginText(72, 730)
    text_object.setFont("Helvetica", 10)

    # Adding the suggestions as text in the new page
    for line in suggestion_text.split('\n'):
        text_object.textLine(line)

    c.drawText(text_object)
    c.showPage()
    c.save()

    packet.seek(0)
    new_pdf = PdfReader(packet)
    writer.add_page(new_pdf.pages[0])

    # Write final optimized PDF
    with open(temp_file, "wb") as f:
        writer.write(f)

    return temp_file

def get_download_link(file, label):
    if isinstance(file, BytesIO):
        b64 = base64.b64encode(file.getvalue()).decode()
        return f'<a href="data:file/pdf;base64,{b64}" download="{label}">{label}</a>'
    else:
        with open(file, "rb") as f:
            data = f.read()
        b64 = base64.b64encode(data).decode()
        return f'<a href="data:file/pdf;base64,{b64}" download="{os.path.basename(file)}">{label}</a>'

# --- Streamlit UI ---
st.set_page_config(page_title="AI Resume Optimizer", layout="wide")
st.title("📄 AI Resume Optimizer")
st.markdown("Upload your resume and a job description to get optimization suggestions")

col1, col2 = st.columns(2)

with col1:
    uploaded_resume = st.file_uploader("Upload Resume (PDF or DOCX)", type=["pdf", "docx"])
with col2:
    job_file = st.file_uploader("Upload Job Description (PDF or DOCX)", type=["pdf", "docx"])

if st.button("Optimize Resume") and uploaded_resume and job_file:
    with st.spinner("Analyzing your resume..."):
        resume_text = extract_text_from_pdf(uploaded_resume) if uploaded_resume.name.endswith(".pdf") else extract_text_from_docx(uploaded_resume)
        job_desc = extract_text_from_pdf(job_file) if job_file.name.endswith(".pdf") else extract_text_from_docx(job_file)

        suggestions = suggest_improvements(resume_text, job_desc)

        os.makedirs("optimized", exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        original_path = f"optimized/original_{timestamp}.txt"
        optimized_path = f"optimized/optimized_{timestamp}.txt"
        diff_path = f"optimized/diff_{timestamp}.txt"

        with open(original_path, "w", encoding="utf-8") as f:
            f.write(resume_text)
        with open(optimized_path, "w", encoding="utf-8") as f:
            f.write(resume_text)
        with open(diff_path, "w", encoding="utf-8") as f:
            f.write(generate_diff(resume_text, resume_text))

        optimized_docx_path = None
        optimized_pdf = None

        if uploaded_resume.name.endswith(".docx"):
            optimized_docx_path = enhance_docx(uploaded_resume, suggestions)
        elif uploaded_resume.name.endswith(".pdf"):
            optimized_pdf = enhance_pdf(uploaded_resume, suggestions)

    st.success("Analysis Complete!")

    st.subheader("Optimization Suggestions")
    for suggestion in suggestions:
        with st.expander(f"{suggestion['category']} - {suggestion['priority']}"):
            st.markdown(f"{suggestion['suggestion']}")

    st.subheader("Download Results")
    if optimized_docx_path:
        st.markdown(get_download_link(optimized_docx_path, "Download Optimized Resume (.docx)"), unsafe_allow_html=True)
    if optimized_pdf:
        st.markdown(get_download_link(optimized_pdf, "Download Optimized Resume (PDF)"), unsafe_allow_html=True)
    st.markdown(get_download_link(optimized_path, "Download Optimized Resume (Text)"), unsafe_allow_html=True)
    st.markdown(get_download_link(diff_path, "Download Changes"), unsafe_allow_html=True)

    st.subheader("Changes Preview")
    st.text_area("Diff View", value=generate_diff(resume_text, resume_text), height=300)

    st.subheader("Extracted Sections")
    for section, content in extract_sections(resume_text).items():
        with st.expander(section.capitalize()):
            st.text(content)

    st.subheader("Named Entities Found")
    for ent_text, ent_label in extract_entities(resume_text):
        st.markdown(f"- {ent_label}: {ent_text}")